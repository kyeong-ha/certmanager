import glob, os, re, shutil, subprocess, tempfile
from django.conf import settings
from django.core.files.base import ContentFile
from typing import Dict, List
from docx import Document
from docx.shared import Cm
from docx.oxml import parse_xml
from docx.text.run import Run

from .licenses import LICENSES

# ────────────────────────────
# 1) ODT → DOCX (임시 폴더)
# ────────────────────────────
def _odt_to_docx(odt_path: str) -> str:
    tmp_dir = tempfile.mkdtemp()
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmp_dir, odt_path],
        check=True
    )
    return os.path.join(
        tmp_dir,
        os.path.splitext(os.path.basename(odt_path))[0] + ".docx",
    )


# ────────────────────────────
# 2) 플레이스홀더 치환
#    (run 파편·공백·제로폭 무시)
# ────────────────────────────
_ZWSP = "\u200B\u200C\u200D\uFEFF"
CLEAN = re.compile(rf"[ \t\r\n{_ZWSP}]")

def _replace_runs(runs: List[Run], mapping: Dict[str, str]):
    collecting, buf, buf_runs = False, "", []
    for run in runs:
        for ch in run.text or "":
            if not collecting:
                if ch == "{":
                    collecting, buf, buf_runs = True, "{", [run]
            else:
                buf += ch
                buf_runs.append(run)
                if ch == "}":
                    key = CLEAN.sub("", buf)          # {성명} 형태
                    if key in mapping:
                        buf_runs[0].text = mapping[key]
                        for r in buf_runs[1:]:
                            r.text = ""
                    collecting, buf, buf_runs = False, "", []

def _replace_placeholders(doc: Document, mapping: Dict[str, str]):
    def _paras(d: Document):
        for p in d.paragraphs:
            yield p
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        yield p
    for para in _paras(doc):
        _replace_runs(para.runs, mapping)


# ────────────────────────────
# 3) 사진 삽입 (3 × 4 cm)
# ────────────────────────────
def _insert_photo(doc: Document, img_path: str):
    if not (img_path and os.path.isfile(img_path)):
        return
    p = doc.add_paragraph()
    r = p.add_run()
    pic = r.add_picture(img_path, width=Cm(2.87), height=Cm(3.82))
    inline = pic._inline
    anchor = f"""
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        simplePos="0" relativeHeight="0" behindDoc="1" allowOverlap="1">
        <wp:positionH relativeFrom="column"><wp:posOffset>4480000</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>3265200</wp:posOffset></wp:positionV>
        <wp:extent cx="1080000" cy="1440000"/><wp:wrapNone/>
        <wp:docPr id="1" name="Picture1"/>{inline.graphic.xml}
    </wp:anchor>"""
    inline.getparent().replace(inline, parse_xml(anchor))


# ────────────────────────────
# 4) DOCX → PDF
#    tmp 변환 후 base.pdf 하나만 이동
# ────────────────────────────
def _docx_to_pdf(docx_path: str, final_dir: str, base: str) -> str:
    """
    DOCX → PDF 
    - /tmp 디렉터리에서 변환
    - output_dir/{base}.pdf로 파일이동(덮어쓰기)
    """
    
    # 1. tmp 폴더에서 PDF 생성
    tmp_pdf_dir = tempfile.mkdtemp()
    subprocess.run(
        [ "libreoffice", "--headless", "--convert-to", 
          "pdf:writer_pdf_Export:DefaultAsianParagraphSpacing=false,EmbedStandardFonts=true",
          "--outdir", tmp_pdf_dir,
        docx_path ],
        check=True,
    )
    produced = next(p for p in os.listdir(tmp_pdf_dir) if p.endswith(".pdf"))
    produced_path = os.path.join(tmp_pdf_dir, produced)
    dst = os.path.join(final_dir, f"{base}.pdf")

    # 2. 결과 PDF 이동 
    shutil.move(produced_path, dst)
    shutil.rmtree(tmp_pdf_dir, ignore_errors=True)

    # 3. 변환 후: 같은 issue 패턴 PDF 중 dst 를 제외하고 모두 삭제
    for old in glob.glob(os.path.join(final_dir, f"{base}*.pdf")):
        if old != dst:
            try:
                os.remove(old)
            except OSError:
                pass

    return dst


# ────────────────────────────
# 5) 
# - ODT 템플릿을 LibreOffice로 .docx로 변환 → _odt_to_docx
# - 치환할 텍스트 입력 → _replace_placeholders
# - 사진 삽입 → _insert_photo
# - .docx → .pdf 변환 → _docx_to_pdf
# ────────────────────────────
def generate_certificate_pdf(cert, output_dir: str) -> str:
    """
    PDF 한 개만 (issue_number.pdf) output_dir 에 남긴다.
    """
    issue_number = cert.issue_number.replace('/', '-')
    pdf_path = os.path.join(output_dir, f"{issue_number}.pdf")
    
    # 템플릿 → DOCX (임시)
    odt_tpl = os.path.join(settings.BASE_DIR, ".local/staticfiles/templates", "phoenix_ver1.odt")
    docx_tmp = _odt_to_docx(odt_tpl)

    # DOCX 로드 & 텍스트 치환
    doc = Document(docx_tmp)
    replacements = {
        "{성명}": cert.user.user_name,
        "{생년월일}": cert.user.birth_date.strftime("%Y.%m.%d") if cert.user.birth_date else "",
        "{자격과정}": cert.course_name,
        "{발급번호}": cert.issue_number,
        "{등록번호}": LICENSES.get(cert.course_name) or "",
        "{발급일자}": cert.issue_date.strftime("%Y년 %m월 %d일") if cert.issue_date else "",
    }
    _replace_placeholders(doc, replacements)

    # 사진
    if cert.user.photo:
        _insert_photo(doc, cert.user.photo.path)

    # DOCX 저장 (tmp) & PDF 변환
    tmp_dir = tempfile.mkdtemp()
    word_path = os.path.join(tmp_dir, f"{issue_number}.docx")
    doc.save(word_path)
    pdf_path = _docx_to_pdf(word_path, output_dir, issue_number)

    # 임시 DOCX & 템플릿 복사본 정리
    for p in (word_path, docx_tmp):
        try:
            os.remove(p)
        except OSError:
            pass
    shutil.rmtree(tmp_dir, ignore_errors=True)
    
    return pdf_path


# ────────────────────────────
# 6) 파일명을 {issue_number}.pdf 에 맞춰서 저장
# ────────────────────────────
def save_certificate_pdf(certificate, pdf_path: str):
    from api.cert.models.Certificate import Certificate  # 순환참조 방지

    filename = f"{certificate.issue_number}.pdf"
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 파일이 존재하지 않음: {pdf_path}")

    with open(pdf_path, "rb") as fp:
        certificate.copy_file.save(filename, ContentFile(fp.read()), save=True)